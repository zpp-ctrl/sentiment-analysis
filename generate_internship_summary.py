# -*- coding: utf-8 -*-
"""
============================================================
生成实习总结三件套：图表(PNG) + Word(.docx) + Excel(.xlsx)
============================================================
"""
import os
import sys
from datetime import date

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, PROJECT_ROOT)

from config import OUTPUT_DIR

CHART_DIR = os.path.join(OUTPUT_DIR, "charts")
os.makedirs(CHART_DIR, exist_ok=True)

DOCX_PATH = os.path.join(OUTPUT_DIR, "实习总结_财经舆情预测系统.docx")
XLSX_PATH = os.path.join(OUTPUT_DIR, "实习总结_财经舆情预测系统.xlsx")

# ==================== 主题色板（品牌中性，色盲安全） ====================
C_BLUE = "#2a78d6"
C_RED = "#e34948"
C_GRAY = "#898781"
C_GRAY_LIGHT = "#e1e0d9"
C_INK = "#0b0b0b"
C_INK2 = "#52514e"
C_GREEN = "#1baf7a"
C_VIOLET = "#4a3aa7"
C_ORANGE = "#eb6834"

# ==================== 核心数据 ====================

TECH_STACK = [
    ("开发语言", "Python 3.10"),
    ("数据采集", "Playwright 浏览器自动化 + asyncio 并发"),
    ("反爬处理", "登录态 Cookie 管理、断点续存、请求头修正"),
    ("情感分析", "DeepSeek API（LLM 大模型）+ 关键词词典兜底"),
    ("视频/图像", "Whisper 语音转写 + OpenCV 抽帧 + Vision API / EasyOCR"),
    ("数据处理", "pandas / numpy"),
    ("数据存储", "MySQL + SQLite 自动回退"),
    ("报表输出", "openpyxl（Excel）"),
    ("任务调度", "Windows 任务计划程序 / Linux cron"),
    ("消息推送", "Server酱（微信）"),
]

MODULES = [
    ("module1", "抖音账号采集", "主页 XHR 采集 · 多线程 · 断点续存"),
    ("module1.5", "视频内容提取", "Whisper 转写 + 视觉分析"),
    ("module2", "文本情感分类", "DeepSeek LLM 批量多空分类"),
    ("module3", "情绪指标汇总", "按博主等级统计看多/看空/中性占比"),
    ("module4", "指数数据", "akshare 行情 + 交易日历"),
    ("module5", "涨跌预测", "关键词分流 + 粉丝加权 + 置信度加权 + 时间衰减"),
    ("module6", "回测验证", "T+1 真实涨跌比对"),
    ("module7", "数据存储", "MySQL/SQLite + Excel 报表"),
    ("module8", "任务调度", "Windows 计划任务 / cron"),
    ("notify", "消息推送", "Server酱 微信推送"),
]

PROBLEMS = [
    ("抖音搜索被验证码拦截", "关键词搜索永远返回 0 条", "改用博主主页 XHR 采集，跳过搜索"),
    ("主页采集 0 条", "Sec-Fetch-Dest/Mode 请求头被错误加到 XHR 上", "去除多余请求头，恢复每主页 20+ 条"),
    ("视频下载全部失败", "抖音视频需签名 CDN 地址，拿不到", "回退纯文本分析（待解决）"),
    ("LLM 余额不足(402)", "约 40% 帖子被降级为中性，预测失真", "充值后补跑全量分类"),
    ("DB 写入 NaN 报错", "int(None) 崩溃，预测未入库", "新增 _safe_int/_safe_float 保护"),
    ("预测目标日期错误", "原代码 target_day 误设为当日", "改为 get_next_trading_day 次日"),
]

PREDICTIONS = [
    ("上证指数", "偏多", 0.6070, 0.553, 0.447),
    ("沪深300", "偏多", 0.6219, 0.561, 0.439),
    ("中证500", "偏多", 0.5631, 0.531, 0.469),
    ("中证1000", "偏多", 0.5796, 0.540, 0.460),
]

# ==================== 图表 ====================
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei"]
plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["figure.dpi"] = 150
plt.rcParams["savefig.facecolor"] = "#ffffff"


def _style_ax(ax):
    for s in ["top", "right"]:
        ax.spines[s].set_visible(False)
    for s in ["left", "bottom"]:
        ax.spines[s].set_color(C_GRAY_LIGHT)
    ax.tick_params(colors=C_INK2, labelsize=9)
    ax.yaxis.grid(True, color=C_GRAY_LIGHT, linewidth=0.8)
    ax.set_axisbelow(True)


def chart1_pipeline():
    """流水线架构图（纵向）"""
    flow = [
        ("采集", "抖音博主主页 · 2491 条/日"),
        ("视频提取", "Whisper 转写 + 视觉分析"),
        ("情感分类", "DeepSeek LLM · 多空"),
        ("情绪汇总", "按博主等级统计"),
        ("涨跌预测", "加权 · 次日方向"),
        ("回测验证", "T+1 真实比对"),
        ("存储报表", "MySQL + Excel"),
        ("推送", "Server酱 微信"),
    ]
    fig, ax = plt.subplots(figsize=(5.2, 7.4))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, len(flow) * 2 + 1)
    ax.axis("off")

    n = len(flow)
    box_w, box_h = 6.2, 1.05
    x = (10 - box_w) / 2
    for i, (name, desc) in enumerate(flow):
        y_top = (n - i) * 2 - 1
        box = FancyBboxPatch(
            (x, y_top - box_h), box_w, box_h,
            boxstyle="round,pad=0.06,rounding_size=0.15",
            linewidth=1.0, edgecolor=C_BLUE, facecolor="#eaf2fc", zorder=3,
        )
        ax.add_patch(box)
        ax.text(x + 0.35, y_top - box_h / 2 + 0.16, f"{i+1}",
                fontsize=12, fontweight="bold", color=C_BLUE,
                va="center", ha="center", zorder=4)
        ax.text(x + 0.85, y_top - box_h / 2 + 0.18, name,
                fontsize=12, fontweight="bold", color=C_INK, va="center", zorder=4)
        ax.text(x + 0.85, y_top - box_h / 2 - 0.20, desc,
                fontsize=8.5, color=C_INK2, va="center", zorder=4)
        if i < n - 1:
            arrow = FancyArrowPatch(
                (x + box_w / 2, y_top - box_h - 0.18),
                (x + box_w / 2, y_top - box_h - 0.85),
                arrowstyle="-|>", mutation_scale=14,
                linewidth=1.4, color=C_BLUE, zorder=2,
            )
            ax.add_patch(arrow)

    ax.set_title("系统流水线架构（端到端闭环）", fontsize=13, fontweight="bold",
                 color=C_INK, pad=12)
    fig.tight_layout()
    p = os.path.join(CHART_DIR, "01_pipeline.png")
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def chart2_neutral_trend():
    """中性占比演进（系统从'失效'到'跑通'）"""
    labels = ["07-31", "08-03", "08-10", "08-18"]
    neutral = [99.02, 99.49, 100.0, 60.94]
    fig, ax = plt.subplots(figsize=(6.4, 3.6))
    ax.plot(labels, neutral, marker="o", linewidth=2.2, color=C_BLUE,
            markersize=7, markerfacecolor="white", markeredgewidth=2,
            markeredgecolor=C_BLUE, zorder=3)
    for x, y in zip(labels, neutral):
        ax.annotate(f"{y:.1f}%", (x, y), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=9, color=C_INK2)
    ax.annotate("首次完整运行\nLLM 全部生效", xy=(3, 60.94),
                xytext=(2.15, 72), fontsize=9, color=C_RED,
                arrowprops=dict(arrowstyle="->", color=C_RED, lw=1.2))
    ax.set_ylim(40, 108)
    ax.set_ylabel("中性帖子占比 (%)", fontsize=9, color=C_INK2)
    _style_ax(ax)
    ax.set_title("中性占比随系统修复显著下降", fontsize=12, fontweight="bold",
                 color=C_INK, pad=10)
    fig.tight_layout()
    p = os.path.join(CHART_DIR, "02_neutral_trend.png")
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def chart3_before_after():
    """余额修复前后情感分布对比"""
    cats = ["看多", "看空", "中性"]
    before = [17.42, 6.66, 75.91]
    after = [24.09, 14.97, 60.94]
    colors = [C_BLUE, C_RED, C_GRAY]
    x = range(len(cats))
    w = 0.38
    fig, ax = plt.subplots(figsize=(6.4, 3.6))
    b1 = ax.bar([i - w / 2 for i in x], before, w, label="修复前(余额不足)",
                color=[_c for _c in colors], alpha=0.55, linewidth=0)
    b2 = ax.bar([i + w / 2 for i in x], after, w, label="修复后(补跑)",
                color=colors, linewidth=0)
    for bars in (b1, b2):
        for r in bars:
            ax.annotate(f"{r.get_height():.1f}%",
                        (r.get_x() + r.get_width() / 2, r.get_height()),
                        textcoords="offset points", xytext=(0, 3),
                        ha="center", fontsize=8.5, color=C_INK2)
    ax.set_xticks(list(x))
    ax.set_xticklabels(cats, fontsize=10)
    ax.set_ylabel("占比 (%)", fontsize=9, color=C_INK2)
    ax.set_ylim(0, 88)
    _style_ax(ax)
    ax.legend(fontsize=8.5, frameon=False, loc="upper right")
    ax.set_title("余额修复后：看空占比翻倍，中性大幅回落", fontsize=12,
                 fontweight="bold", color=C_INK, pad=10)
    fig.tight_layout()
    p = os.path.join(CHART_DIR, "03_before_after.png")
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def chart4_prediction():
    """四大指数次日预测多空得分"""
    names = [p[0] for p in PREDICTIONS][::-1]
    bull = [p[3] for p in PREDICTIONS][::-1]
    bear = [p[4] for p in PREDICTIONS][::-1]
    conf = [p[2] * 100 for p in PREDICTIONS][::-1]
    y = range(len(names))
    h = 0.34
    fig, ax = plt.subplots(figsize=(6.4, 3.6))
    ax.barh([i + h / 2 for i in y], bull, h, color=C_BLUE, label="看多得分", linewidth=0)
    ax.barh([i - h / 2 for i in y], bear, h, color=C_RED, label="看空得分", linewidth=0)
    for i, (b, c) in enumerate(zip(bull, conf)):
        ax.text(0.515, i, f"{b:.3f}", va="center", ha="left", fontsize=8.5, color=C_INK2)
    ax.set_yticks(list(y))
    ax.set_yticklabels(names, fontsize=10)
    ax.set_xlim(0.4, 0.62)
    ax.set_xlabel("加权得分（0.5 为中性分界）", fontsize=9, color=C_INK2)
    _style_ax(ax)
    ax.axvline(0.5, color=C_GRAY, linewidth=1.0, linestyle="--")
    ax.text(0.5, len(y) - 0.25, " 中性 0.5", fontsize=8, color=C_GRAY, va="center")
    ax.legend(fontsize=8.5, frameon=False, loc="lower right")
    ax.set_title("四大指数次日(8/19)预测：多空加权得分", fontsize=12,
                 fontweight="bold", color=C_INK, pad=10)
    fig.tight_layout()
    p = os.path.join(CHART_DIR, "04_prediction.png")
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def build_charts():
    return {
        "pipeline": chart1_pipeline(),
        "neutral": chart2_neutral_trend(),
        "before_after": chart3_before_after(),
        "prediction": chart4_prediction(),
    }


# ==================== Word ====================
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn


def _set_run_font(run, size=10.5, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    from docx.oxml.ns import qn
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def _add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    _set_run_font(run, size={1: 15, 2: 12.5, 3: 11}.get(level, 11),
                  bold=True, color=RGBColor(0x2A, 0x78, 0xD6))
    return p


def _add_para(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    return p


def _add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _set_run_font(run, size=10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            _set_run_font(run, size=9.5)
    return t


def build_docx(charts):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("实习总结\n财经舆情预测系统")
    _set_run_font(r, size=22, bold=True, color=RGBColor(0x0B, 0x0B, 0x0B))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("基于社交媒体情感分析的指数次日涨跌预测 · 2026年8月")
    _set_run_font(r, size=11, color=RGBColor(0x52, 0x51, 0x4E))

    # 一、项目背景与目标
    _add_heading(doc, "一、项目背景与目标", 1)
    _add_para(doc, "股票市场短期涨跌与投资者情绪高度相关，抖音、小红书等平台聚集了大量财经观点。"
                   "本项目旨在构建一条自动化闭环：自动采集财经博主观点 → 情感分析 → 预测四大宽基指数"
                   "（上证、沪深300、中证500、中证1000）次日涨跌方向 → 回测验证准确率 → 生成日报并推送。")
    _add_para(doc, "核心目标：用可验证的方式，把海量非结构化的社交媒体情绪转化为可执行的次日方向判断。",
              bold=True)

    # 二、工作思路
    _add_heading(doc, "二、工作思路", 1)
    _add_para(doc, "整体方法论是「数据 → 情感 → 预测 → 验证」的闭环：")
    _add_para(doc, "1. 数据是根基：先解决「采得到」——反爬、登录态、断点续存、请求头修正；", bold=False)
    _add_para(doc, "2. 情感是核心：用 LLM 大模型做多空分类，比关键词词典更准确；")
    _add_para(doc, "3. 预测要可验证：T+1 收盘回测，用真实涨跌检验预测命中率，而不是自说自话。")

    # 三、实现途径 / 架构
    _add_heading(doc, "三、实现途径与系统架构", 1)
    _add_para(doc, "系统采用模块化流水线，共 10 个环节，每日定时自动运行：")
    doc.add_picture(charts["pipeline"], width=Inches(5.2))
    _add_table(doc, ["模块", "职责", "关键实现"], MODULES)

    # 四、技术栈
    _add_heading(doc, "四、技术栈", 1)
    _add_table(doc, ["技术方向", "具体技术"], TECH_STACK)

    # 五、思考历程：问题与解决
    _add_heading(doc, "五、思考历程：遇到的问题与解决", 1)
    _add_para(doc, "实习过程中最有价值的部分，是把系统从「搭起来」推进到「真正跑通」。"
                   "下表梳理了主要问题及解决过程：")
    _add_table(doc, ["问题", "根因分析", "解决方案"], PROBLEMS)
    _add_para(doc, "其中「LLM 余额不足」问题直接影响了预测质量，修复前后对比如下图：")
    doc.add_picture(charts["before_after"], width=Inches(5.6))
    doc.add_picture(charts["neutral"], width=Inches(5.6))

    # 六、达到的成果
    _add_heading(doc, "六、达到的成果", 1)
    _add_para(doc, "2026-08-18 系统实现首次端到端完整运行，主要成果如下：")
    _add_table(doc, ["指标", "结果"], [
        ("账号池规模", "2543 个财经博主"),
        ("单日采集帖子", "2491 条（24 小时窗口）"),
        ("情感分类", "看多 600 / 看空 373 / 中性 1518"),
        ("LLM 平均置信度", "0.79"),
        ("次日预测", "8/19 四大指数全部「偏多」"),
        ("数据入库", "2491 帖 + 3 情绪统计 + 4 预测记录"),
    ])
    _add_para(doc, "四大指数次日（8/19）预测方向与多空加权得分：")
    _add_table(doc, ["指数", "方向", "置信度", "看多得分", "看空得分"],
               [(p[0], p[1], f"{p[2]*100:.0f}%", f"{p[3]:.3f}", f"{p[4]:.3f}") for p in PREDICTIONS])
    doc.add_picture(charts["prediction"], width=Inches(5.6))
    _add_para(doc, "注：回测准确率尚无长期数据（历史回测任务未跑通），预测命中率需后续交易日持续验证。",
              size=9)

    # 七、不足与展望
    _add_heading(doc, "七、不足与展望", 1)
    for t in [
        "视频内容提取未打通（抖音签名 CDN 地址问题），当前仅用文字描述；",
        "搜索被验证码拦截，只能靠主页采集，数据量与覆盖面受限；",
        "小红书采集暂未启用，目前仅使用抖音单平台数据；",
        "回测验证闭环尚未完整跑通，预测准确率需长期积累；",
        "后续可引入历史情绪时序特征、代理/IP 池，提升采集稳定性与预测精度。",
    ]:
        _add_para(doc, "· " + t)

    doc.save(DOCX_PATH)
    return DOCX_PATH


# ==================== Excel ====================
def build_xlsx(charts):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    header_fill = PatternFill("solid", fgColor="2A78D6")
    header_font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
    cell_font = Font(name="微软雅黑", size=10)
    title_font = Font(name="微软雅黑", size=14, bold=True, color="0B0B0B")
    thin = Side(style="thin", color="E1E0D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def write_table(ws, start_row, headers, rows, widths):
        for j, h in enumerate(headers):
            c = ws.cell(row=start_row, column=j + 1, value=h)
            c.fill = header_fill
            c.font = header_font
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = border
        for i, row in enumerate(rows):
            for j, v in enumerate(row):
                c = ws.cell(row=start_row + 1 + i, column=j + 1, value=v)
                c.font = cell_font
                c.border = border
                c.alignment = Alignment(vertical="center", wrap_text=(j > 0))
        for j, w in enumerate(widths):
            ws.column_dimensions[get_column_letter(j + 1)].width = w
        return start_row + 1 + len(rows)

    # Sheet 1: 项目概览
    ws = wb.active
    ws.title = "项目概览"
    ws["A1"] = "财经舆情预测系统 — 项目概览"
    ws["A1"].font = title_font
    ws.merge_cells("A1:C1")
    rows = [
        ("项目名称", "每日自动财经舆情预测系统"),
        ("项目周期", "2026-07 ~ 2026-08"),
        ("核心目标", "预测四大宽基指数次日涨跌方向，T+1 回测验证"),
        ("数据源", "抖音财经博主（小红书暂未启用）"),
        ("账号池规模", "2543 个财经博主"),
        ("单日采集量", "2491 条帖子（24h 窗口）"),
        ("情感分析", "DeepSeek LLM 大模型"),
        ("四大指数", "上证 / 沪深300 / 中证500 / 中证1000"),
    ]
    write_table(ws, 3, ["维度", "内容", ""], [(a, b, "") for a, b in rows], [16, 52, 4])

    # Sheet 2: 技术栈
    ws2 = wb.create_sheet("技术栈")
    ws2["A1"] = "技术栈"; ws2["A1"].font = title_font
    write_table(ws2, 2, ["技术方向", "具体技术"], TECH_STACK, [18, 56])

    # Sheet 3: 模块架构
    ws3 = wb.create_sheet("模块架构")
    ws3["A1"] = "系统模块"; ws3["A1"].font = title_font
    write_table(ws3, 2, ["模块", "职责", "关键实现"], MODULES, [14, 20, 44])
    img = XLImage(charts["pipeline"])
    img.width = 420
    ws3.add_image(img, f"A{len(MODULES)+5}")

    # Sheet 4: 问题与解决
    ws4 = wb.create_sheet("问题与解决")
    ws4["A1"] = "问题与解决"; ws4["A1"].font = title_font
    write_table(ws4, 2, ["问题", "根因分析", "解决方案"], PROBLEMS, [24, 36, 34])

    # Sheet 5: 成果数据
    ws5 = wb.create_sheet("成果数据")
    ws5["A1"] = "2026-08-18 首次完整运行成果"; ws5["A1"].font = title_font
    write_table(ws5, 2, ["指标", "结果"], [
        ("账号池规模", "2543 个财经博主"),
        ("单日采集帖子", "2491 条"),
        ("情感分类", "看多 600 / 看空 373 / 中性 1518"),
        ("LLM 平均置信度", "0.79"),
        ("次日(8/19)预测", "四大指数全部「偏多」"),
        ("数据入库", "2491 帖 + 3 统计 + 4 预测"),
    ], [18, 44])
    write_table(ws5, 10, ["指数", "方向", "置信度", "看多得分", "看空得分"],
                [(p[0], p[1], f"{p[2]*100:.0f}%", f"{p[3]:.3f}", f"{p[4]:.3f}") for p in PREDICTIONS],
                [14, 10, 10, 12, 12])
    img2 = XLImage(charts["prediction"])
    img2.width = 430
    ws5.add_image(img2, "A17")
    img3 = XLImage(charts["before_after"])
    img3.width = 430
    ws5.add_image(img3, "A33")

    wb.save(XLSX_PATH)
    return XLSX_PATH


if __name__ == "__main__":
    print("生成图表...")
    charts = build_charts()
    for k, v in charts.items():
        print("  chart:", v)
    print("生成 Word...")
    print("  ", build_docx(charts))
    print("完成 ✅")
